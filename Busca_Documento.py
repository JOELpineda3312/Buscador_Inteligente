import os
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from tkinter.font import Font
import customtkinter as ctk
from PIL import Image, ImageTk
from PyPDF2 import PdfReader
from docx import Document
import openpyxl
import threading
import subprocess
import platform


class AIFileSearchApp:
    def __init__(self, root):
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")

        self.root = root
        self.root.title("Nexus File Intelligence")
        self.root.geometry("1300x900")

        self.title_font = ctk.CTkFont(family="Roboto", size=24, weight="bold")
        self.subtitle_font = ctk.CTkFont(family="Roboto", size=16)

        self.folder_contexts = {
            "Desktop": {"path": os.path.expanduser("~/Desktop"), "priority": 1},
            "Documents": {"path": os.path.expanduser("~/Documents"), "priority": 2},
            "Pictures": {"path": os.path.expanduser("~/Pictures"), "priority": 3},
            "Videos": {"path": os.path.expanduser("~/Videos"), "priority": 4},
            "Music": {"path": os.path.expanduser("~/Music"), "priority": 5},
            "System Root": {"path": "C:\\", "priority": 6},
        }

        self.folder_vars = {name: ctk.BooleanVar(value=False) for name in self.folder_contexts}
        self.search_history = []

        self.create_ui()
        self.results_table.bind("<Double-1>", self.open_file_location)
        self.results_table.bind("<Button-3>", self.show_context_menu)

    def create_ui(self):
        self.main_frame = ctk.CTkFrame(self.root, corner_radius=10)
        self.main_frame.pack(padx=20, pady=20, fill="both", expand=True)

        self.title_label = ctk.CTkLabel(
            self.main_frame, text="Nexus File Intelligence", font=self.title_font
        )
        self.title_label.pack(pady=(20, 10))

        search_frame = ctk.CTkFrame(self.main_frame, corner_radius=10)
        search_frame.pack(pady=10, padx=20, fill="x")

        self.search_entry = ctk.CTkEntry(
            search_frame,
            placeholder_text="🔍 Enter search term...",
            width=600,
            font=self.subtitle_font,
        )
        self.search_entry.pack(side="left", padx=10, pady=10, expand=True)

        folder_frame = ctk.CTkFrame(self.main_frame, corner_radius=10)
        folder_frame.pack(pady=10, padx=20, fill="x")

        for name, context in self.folder_contexts.items():
            checkbox = ctk.CTkCheckBox(
                folder_frame,
                text=f"{name} (Priority {context['priority']})",
                variable=self.folder_vars[name],
                font=self.subtitle_font,
            )
            checkbox.pack(side="left", padx=10)

        self.search_button = ctk.CTkButton(
            self.main_frame, text="Initiate Search", command=self.start_search, font=self.subtitle_font
        )
        self.search_button.pack(pady=10)

        self.progress_bar = ctk.CTkProgressBar(self.main_frame)
        self.progress_bar.pack(pady=10, padx=20, fill="x")
        self.progress_bar.set(0)

        self.results_table = ttk.Treeview(
            self.main_frame, columns=("Filename", "Path", "Type", "Relevance"), show="headings"
        )
        for col in ("Filename", "Path", "Type", "Relevance"):
            self.results_table.heading(col, text=col)
        self.results_table.pack(pady=10, padx=20, fill="both", expand=True)

        self.preview_frame = ctk.CTkFrame(self.main_frame, corner_radius=10)
        self.preview_frame.pack(pady=10, padx=20, fill="x")

        self.preview_label = ctk.CTkLabel(self.preview_frame, text="File Preview", font=self.subtitle_font)
        self.preview_label.pack(pady=10)

    def start_search(self):
        search_term = self.search_entry.get().strip()
        if not search_term:
            messagebox.showerror("Error", "Por favor, introduce un texto para buscar.")
            return

        selected_folders = [
            context["path"] for name, context in self.folder_contexts.items() if self.folder_vars[name].get()
        ]
        if not selected_folders:
            messagebox.showerror("Error", "Por favor, selecciona al menos una carpeta para buscar.")
            return

        self.progress_bar.set(0)
        for i in self.results_table.get_children():
            self.results_table.delete(i)

        threading.Thread(target=self.search_files, args=(search_term, selected_folders), daemon=True).start()

    def search_files(self, search_term, folders):
        results = []
        for folder in folders:
            for root, _, files in os.walk(folder):
                for file in files:
                    if search_term.lower() in file.lower():
                        full_path = os.path.join(root, file)
                        results.append(
                            (file, full_path, os.path.splitext(file)[1], self.calculate_relevance(search_term, file))
                        )

        self.root.after(0, self.update_results, results, search_term)

    def calculate_relevance(self, search_term, file_name):
        return file_name.lower().count(search_term.lower())

    def update_results(self, results, search_term):
        results.sort(key=lambda x: x[3], reverse=True)
        if not results:
            messagebox.showinfo("Búsqueda", f"No se encontraron archivos con el nombre '{search_term}'.")
        for result in results:
            self.results_table.insert("", "end", values=result)
        self.progress_bar.set(1)

    def open_file_location(self, event):
        selected_item = self.results_table.selection()
        if not selected_item:
            return

        file_path = self.results_table.item(selected_item, "values")[1]
        file_path = os.path.normpath(file_path)

        try:
            if platform.system() == "Windows":
                os.startfile(os.path.dirname(file_path))
            elif platform.system() == "Darwin":
                subprocess.run(["open", "-R", file_path], check=True)
            else:
                subprocess.run(["xdg-open", os.path.dirname(file_path)], check=True)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir la ubicación del archivo: {str(e)}")

    def show_context_menu(self, event):
        context_menu = tk.Menu(self.root, tearoff=0)
        context_menu.add_command(label="Abrir Archivo", command=self.open_file)
        context_menu.add_command(label="Previsualizar Archivo", command=self.preview_file)
        context_menu.tk_popup(event.x_root, event.y_root)

    def open_file(self):
        selected_item = self.results_table.selection()
        if not selected_item:
            messagebox.showerror("Error", "Por favor, selecciona un archivo.")
            return

        file_path = self.results_table.item(selected_item, "values")[1]
        try:
            if platform.system() == "Windows":
                os.startfile(file_path)
            elif platform.system() == "Darwin":
                subprocess.run(["open", file_path], check=True)
            else:
                subprocess.run(["xdg-open", file_path], check=True)
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo abrir el archivo: {str(e)}")

    def preview_file(self):
        selected_item = self.results_table.selection()
        if not selected_item:
            messagebox.showerror("Error", "Por favor, selecciona un archivo.")
            return

        file_path = self.results_table.item(selected_item, "values")[1]
        try:
            if file_path.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp")):
                self.preview_image(file_path)
            elif file_path.lower().endswith(".pdf"):
                self.preview_pdf(file_path)
            elif file_path.lower().endswith(".docx"):
                self.preview_docx(file_path)
            elif file_path.lower().endswith(".xlsx"):
                self.preview_xlsx(file_path)
            else:
                messagebox.showinfo("Previsualización", "No se puede previsualizar este tipo de archivo.")
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo previsualizar el archivo: {str(e)}")

    def preview_image(self, file_path):
        img = Image.open(file_path)
        img.thumbnail((400, 400))
        img_tk = ImageTk.PhotoImage(img)
        self.preview_label.config(image=img_tk)
        self.preview_label.image = img_tk

    def preview_pdf(self, file_path):
        reader = PdfReader(file_path)
        first_page = reader.pages[0]
        text = first_page.extract_text()
        self.preview_label.config(text=text)

    def preview_docx(self, file_path):
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        self.preview_label.config(text=text)
    def preview_docx(self, file_path):
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        self.preview_label.config(text=text)
    
    def preview_xlsx(self, file_path):
        wb = openpyxl.load_workbook(file_path)
        sheet = wb.active
        data = "\n".join([", ".join([str(cell.value) for cell in row]) for row in sheet.iter_rows()])
        self.preview_label.config(text=data)

def main():
    root = ctk.CTk()
    app = AIFileSearchApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
